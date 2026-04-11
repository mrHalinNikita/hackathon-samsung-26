from pathlib import Path
from docx import Document

from src.parsers.base import BaseParser, ParsedContent


class DocxParser(BaseParser):
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]
    
    async def parse(self, filepath: Path) -> ParsedContent:
        text_parts = []
        metadata = {"path": str(filepath)}
        
        try:
            doc = Document(filepath)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            if doc.core_properties:
                metadata.update({
                    "author": doc.core_properties.author,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                })
        
        except Exception as e:
            return ParsedContent(
                text="",
                metadata=metadata,
                errors=[f"DOCX parse error: {e}"],
            )
        
        full_text = "\n\n".join(text_parts)
        
        return ParsedContent(
            text=full_text,
            metadata=metadata,
            word_count=len(full_text.split()),
            char_count=len(full_text),
        )